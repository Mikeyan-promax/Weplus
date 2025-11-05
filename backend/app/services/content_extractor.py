"""
WePlus 后台管理系统 - 文件内容提取服务
支持PDF文本提取、图片OCR、文档格式转换等功能
"""

import os
import io
import logging
from typing import Dict, List, Optional, Tuple, Any
from pathlib import Path
import hashlib
import mimetypes
from datetime import datetime

# 文档处理库
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PDF处理库未安装，PDF功能将不可用")

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR库未安装，图片文字识别功能将不可用")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("DOCX处理库未安装，Word文档功能将不可用")

try:
    import openpyxl
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logging.warning("Excel处理库未安装，Excel文档功能将不可用")

# 配置日志
logger = logging.getLogger(__name__)

class ContentExtractionResult:
    """内容提取结果类"""
    
    def __init__(self):
        self.success: bool = False
        self.content: str = ""
        self.metadata: Dict[str, Any] = {}
        self.error: Optional[str] = None
        self.file_type: str = ""
        self.page_count: int = 0
        self.word_count: int = 0
        self.extraction_time: float = 0.0
        self.confidence: float = 0.0  # OCR置信度
        
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式"""
        return {
            "success": self.success,
            "content": self.content,
            "metadata": self.metadata,
            "error": self.error,
            "file_type": self.file_type,
            "page_count": self.page_count,
            "word_count": self.word_count,
            "extraction_time": self.extraction_time,
            "confidence": self.confidence
        }

class ContentExtractor:
    """文件内容提取器"""
    
    def __init__(self):
        self.supported_types = {
            'pdf': ['.pdf'],
            'image': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif'],
            'text': ['.txt', '.md', '.csv'],
            'word': ['.docx', '.doc'],
            'excel': ['.xlsx', '.xls'],
            'powerpoint': ['.pptx', '.ppt']
        }
        
        # OCR配置
        self.ocr_config = {
            'lang': 'chi_sim+eng',  # 中文简体 + 英文
            'config': '--oem 3 --psm 6'  # OCR引擎模式和页面分割模式
        }
        
        # 文本清理配置
        self.text_cleanup_config = {
            'min_line_length': 3,
            'remove_empty_lines': True,
            'normalize_whitespace': True,
            'remove_special_chars': False
        }
    
    def get_file_type(self, file_path: str) -> str:
        """获取文件类型"""
        file_ext = Path(file_path).suffix.lower()
        
        for file_type, extensions in self.supported_types.items():
            if file_ext in extensions:
                return file_type
        
        return 'unknown'
    
    def is_supported(self, file_path: str) -> bool:
        """检查文件是否支持提取"""
        return self.get_file_type(file_path) != 'unknown'
    
    def extract_content(self, file_path: str) -> ContentExtractionResult:
        """提取文件内容"""
        start_time = datetime.now()
        result = ContentExtractionResult()
        
        try:
            if not os.path.exists(file_path):
                result.error = "文件不存在"
                return result
            
            file_type = self.get_file_type(file_path)
            result.file_type = file_type
            
            # 根据文件类型选择提取方法
            if file_type == 'pdf':
                result = self._extract_pdf_content(file_path, result)
            elif file_type == 'image':
                result = self._extract_image_content(file_path, result)
            elif file_type == 'text':
                result = self._extract_text_content(file_path, result)
            elif file_type == 'word':
                result = self._extract_word_content(file_path, result)
            elif file_type == 'excel':
                result = self._extract_excel_content(file_path, result)
            else:
                result.error = f"不支持的文件类型: {file_type}"
                return result
            
            # 计算提取时间
            end_time = datetime.now()
            result.extraction_time = (end_time - start_time).total_seconds()
            
            # 后处理
            if result.success and result.content:
                result.content = self._clean_text(result.content)
                result.word_count = len(result.content.split())
                
                # 添加基本元数据
                result.metadata.update({
                    'file_size': os.path.getsize(file_path),
                    'file_name': os.path.basename(file_path),
                    'extraction_method': file_type,
                    'content_hash': hashlib.md5(result.content.encode()).hexdigest()
                })
            
            logger.info(f"内容提取完成: {file_path}, 类型: {file_type}, 成功: {result.success}")
            
        except Exception as e:
            result.error = f"内容提取失败: {str(e)}"
            logger.error(f"内容提取异常: {file_path}, 错误: {e}")
        
        return result
    
    def _extract_pdf_content(self, file_path: str, result: ContentExtractionResult) -> ContentExtractionResult:
        """提取PDF内容"""
        if not PDF_AVAILABLE:
            result.error = "PDF处理库未安装"
            return result
        
        try:
            content_parts = []
            
            # 尝试使用pdfplumber（更好的文本提取）
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    result.page_count = len(pdf.pages)
                    
                    for page_num, page in enumerate(pdf.pages, 1):
                        text = page.extract_text()
                        if text:
                            content_parts.append(f"=== 第{page_num}页 ===\n{text}\n")
                        
                        # 提取表格
                        tables = page.extract_tables()
                        for table_num, table in enumerate(tables, 1):
                            if table:
                                table_text = self._format_table(table)
                                content_parts.append(f"=== 第{page_num}页 表格{table_num} ===\n{table_text}\n")
                
                result.metadata['extraction_method'] = 'pdfplumber'
                
            except Exception as e:
                logger.warning(f"pdfplumber提取失败，尝试PyPDF2: {e}")
                
                # 备用方案：使用PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    result.page_count = len(pdf_reader.pages)
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        text = page.extract_text()
                        if text:
                            content_parts.append(f"=== 第{page_num}页 ===\n{text}\n")
                
                result.metadata['extraction_method'] = 'PyPDF2'
            
            if content_parts:
                result.content = "\n".join(content_parts)
                result.success = True
            else:
                result.error = "PDF中未找到可提取的文本内容"
            
        except Exception as e:
            result.error = f"PDF提取失败: {str(e)}"
        
        return result
    
    def _extract_image_content(self, file_path: str, result: ContentExtractionResult) -> ContentExtractionResult:
        """提取图片中的文字（OCR）"""
        if not OCR_AVAILABLE:
            result.error = "OCR库未安装"
            return result
        
        try:
            # 打开图片
            with Image.open(file_path) as image:
                # 图片预处理
                processed_image = self._preprocess_image(image)
                
                # OCR识别
                text = pytesseract.image_to_string(
                    processed_image,
                    lang=self.ocr_config['lang'],
                    config=self.ocr_config['config']
                )
                
                # 获取置信度信息
                try:
                    data = pytesseract.image_to_data(
                        processed_image,
                        lang=self.ocr_config['lang'],
                        output_type=pytesseract.Output.DICT
                    )
                    
                    confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
                    if confidences:
                        result.confidence = sum(confidences) / len(confidences) / 100.0
                    
                except Exception as e:
                    logger.warning(f"获取OCR置信度失败: {e}")
                    result.confidence = 0.5  # 默认置信度
                
                if text.strip():
                    result.content = text
                    result.success = True
                    result.metadata.update({
                        'image_size': image.size,
                        'image_mode': image.mode,
                        'ocr_lang': self.ocr_config['lang']
                    })
                else:
                    result.error = "图片中未识别到文字内容"
        
        except Exception as e:
            result.error = f"图片OCR失败: {str(e)}"
        
        return result
    
    def _extract_text_content(self, file_path: str, result: ContentExtractionResult) -> ContentExtractionResult:
        """提取纯文本内容"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'ascii']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        result.content = content
                        result.success = True
                        result.metadata['encoding'] = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if not result.success:
                result.error = "无法识别文件编码"
        
        except Exception as e:
            result.error = f"文本提取失败: {str(e)}"
        
        return result
    
    def _extract_word_content(self, file_path: str, result: ContentExtractionResult) -> ContentExtractionResult:
        """提取Word文档内容"""
        if not DOCX_AVAILABLE:
            result.error = "Word文档处理库未安装"
            return result
        
        try:
            doc = Document(file_path)
            content_parts = []
            
            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            # 提取表格
            for table_num, table in enumerate(doc.tables, 1):
                table_text = self._extract_word_table(table)
                if table_text:
                    content_parts.append(f"=== 表格{table_num} ===\n{table_text}")
            
            if content_parts:
                result.content = "\n\n".join(content_parts)
                result.success = True
                result.metadata.update({
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables)
                })
            else:
                result.error = "Word文档中未找到内容"
        
        except Exception as e:
            result.error = f"Word文档提取失败: {str(e)}"
        
        return result
    
    def _extract_excel_content(self, file_path: str, result: ContentExtractionResult) -> ContentExtractionResult:
        """提取Excel内容"""
        if not EXCEL_AVAILABLE:
            result.error = "Excel处理库未安装"
            return result
        
        try:
            workbook = load_workbook(file_path, read_only=True)
            content_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_content = []
                
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_data):  # 跳过空行
                        sheet_content.append("\t".join(row_data))
                
                if sheet_content:
                    content_parts.append(f"=== 工作表: {sheet_name} ===\n" + "\n".join(sheet_content))
            
            if content_parts:
                result.content = "\n\n".join(content_parts)
                result.success = True
                result.metadata.update({
                    'sheet_count': len(workbook.sheetnames),
                    'sheet_names': workbook.sheetnames
                })
            else:
                result.error = "Excel文件中未找到内容"
        
        except Exception as e:
            result.error = f"Excel提取失败: {str(e)}"
        
        return result
    
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """图片预处理以提高OCR准确率"""
        try:
            # 转换为灰度图
            if image.mode != 'L':
                image = image.convert('L')
            
            # 调整大小（如果图片太小）
            width, height = image.size
            if width < 300 or height < 300:
                scale_factor = max(300 / width, 300 / height)
                new_size = (int(width * scale_factor), int(height * scale_factor))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
            
            return image
            
        except Exception as e:
            logger.warning(f"图片预处理失败: {e}")
            return image
    
    def _format_table(self, table: List[List[str]]) -> str:
        """格式化表格数据"""
        if not table:
            return ""
        
        formatted_rows = []
        for row in table:
            if row:
                # 清理单元格数据
                cleaned_row = [str(cell).strip() if cell else "" for cell in row]
                formatted_rows.append("\t".join(cleaned_row))
        
        return "\n".join(formatted_rows)
    
    def _extract_word_table(self, table) -> str:
        """提取Word表格内容"""
        try:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append("\t".join(row_data))
            
            return "\n".join(table_data)
        except Exception as e:
            logger.warning(f"Word表格提取失败: {e}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """清理提取的文本"""
        if not text:
            return ""
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            
            # 跳过过短的行
            if len(line) < self.text_cleanup_config['min_line_length']:
                if not self.text_cleanup_config['remove_empty_lines']:
                    cleaned_lines.append(line)
                continue
            
            # 标准化空白字符
            if self.text_cleanup_config['normalize_whitespace']:
                line = ' '.join(line.split())
            
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def batch_extract(self, file_paths: List[str]) -> Dict[str, ContentExtractionResult]:
        """批量提取文件内容"""
        results = {}
        
        for file_path in file_paths:
            try:
                result = self.extract_content(file_path)
                results[file_path] = result
            except Exception as e:
                error_result = ContentExtractionResult()
                error_result.error = f"批量提取失败: {str(e)}"
                results[file_path] = error_result
        
        return results
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """获取支持的文件格式"""
        return self.supported_types.copy()
    
    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """验证文件是否可以处理"""
        if not os.path.exists(file_path):
            return False, "文件不存在"
        
        if not os.path.isfile(file_path):
            return False, "不是有效的文件"
        
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            return False, "文件为空"
        
        if file_size > 100 * 1024 * 1024:  # 100MB限制
            return False, "文件过大（超过100MB）"
        
        if not self.is_supported(file_path):
            return False, f"不支持的文件类型: {Path(file_path).suffix}"
        
        return True, "文件验证通过"

# 创建全局实例
content_extractor = ContentExtractor()

# 便捷函数
def extract_file_content(file_path: str) -> ContentExtractionResult:
    """提取文件内容的便捷函数"""
    return content_extractor.extract_content(file_path)

def is_file_supported(file_path: str) -> bool:
    """检查文件是否支持的便捷函数"""
    return content_extractor.is_supported(file_path)

def get_supported_formats() -> Dict[str, List[str]]:
    """获取支持格式的便捷函数"""
    return content_extractor.get_supported_formats()