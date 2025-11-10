"""
RAG系统API路由
提供聊天、文档处理、健康检查等端点
"""
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import logging
from datetime import datetime
import uuid
import os
import tempfile
import json
import asyncio

from ..models.rag_models import (
    DocumentProcessRequest, 
    DocumentProcessResponse,
    QueryRequest, 
    QueryResponse,
    DocumentChunk
)
from ..services.rag_service import RAGService
from ..services.postgresql_vector_service import PostgreSQLVectorService
from ..services.document_service import document_service

# 配置日志
logger = logging.getLogger(__name__)

# 创建路由器
router = APIRouter(prefix="/api/rag", tags=["RAG System"])

# 初始化服务
rag_service = RAGService()
postgresql_vector_service = PostgreSQLVectorService()

# 请求模型定义
class ChatRequest(BaseModel):
    """聊天请求模型"""
    message: str
    conversation_history: Optional[List[Dict[str, str]]] = []
    use_rag: bool = True

class ChatResponse(BaseModel):
    """聊天响应模型"""
    response: str
    timestamp: str
    token_usage: Optional[Dict[str, int]] = None
    sources_used: Optional[List[Dict[str, Any]]] = None

class DocumentProcessRequest(BaseModel):
    """文档处理请求模型"""
    content: str
    title: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None

class DocumentProcessResponse(BaseModel):
    """文档处理响应模型"""
    success: bool
    message: str
    document_id: Optional[str] = None
    chunk_count: Optional[int] = None
    processed_at: str

# 临时存储处理过的文档（生产环境应使用数据库）
processed_documents: Dict[str, Dict[str, Any]] = {}

@router.post("/chat/stream")
async def stream_chat_endpoint(request: ChatRequest):
    """
    流式聊天端点
    使用Server-Sent Events返回DeepSeek的实时流式响应
    """
    try:
        logger.info(f"收到流式聊天请求: {request.message[:50]}...")
        
        async def generate_stream():
            try:
                if request.use_rag:
                    # RAG模式：基于文档内容回答
                    logger.info("使用RAG模式进行流式回答")
                    
                    # 从数据库获取所有文档块
                    try:
                        chunks_data = await postgresql_vector_service.get_document_chunks()
                        logger.info(f"从数据库获取到 {len(chunks_data)} 个文档块")
                        
                        if chunks_data:
                            # 检索相关文档块
                            relevant_chunks = await rag_service.retrieve_relevant_chunks_from_db(
                                request.message, chunks_data
                            )
                            
                            # 构建消息
                            context = "\n\n".join([chunk["content"] for chunk in relevant_chunks])
                            system_prompt = f"""你是一个智能校园助手，基于以下知识库内容回答用户问题。
如果知识库中没有相关信息，请诚实地说明并提供一般性建议。

知识库内容：
{context}"""
                            
                            messages = [{"role": "system", "content": system_prompt}]
                            
                            # 添加对话历史
                            if request.conversation_history:
                                messages.extend(request.conversation_history[-10:])
                            
                            # 添加当前消息
                            messages.append({"role": "user", "content": request.message})
                            
                            # 调用DeepSeek流式API
                            response = await rag_service.chat_completion(messages, stream=True)
                            
                            # 处理流式响应
                            for chunk in response["response"]:
                                if chunk.choices[0].delta.content:
                                    content = chunk.choices[0].delta.content
                                    yield f"data: {json.dumps({'content': content, 'finished': False})}\n\n"
                            
                            # 发送完成信号
                            yield f"data: {json.dumps({'content': '', 'finished': True})}\n\n"
                            return
                        else:
                            logger.info("数据库中没有找到文档块，使用普通聊天模式")
                    except Exception as e:
                        logger.error(f"获取文档块失败: {str(e)}")
                        logger.info("回退到普通聊天模式")
                
                # 普通聊天模式
                logger.info("使用普通聊天模式进行流式回答")
                messages = [
                    {"role": "system", "content": "你是一个友好的校园智能助手，可以帮助回答各种问题。"}
                ]
                
                # 添加对话历史
                if request.conversation_history:
                    messages.extend(request.conversation_history[-10:])
                
                # 添加当前消息
                messages.append({"role": "user", "content": request.message})
                
                # 调用DeepSeek流式API
                response = await rag_service.chat_completion(messages, stream=True)
                
                # 处理流式响应
                for chunk in response["response"]:
                    if chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        yield f"data: {json.dumps({'content': content, 'finished': False})}\n\n"
                
                # 发送完成信号
                yield f"data: {json.dumps({'content': '', 'finished': True})}\n\n"
                
            except Exception as e:
                logger.error(f"流式聊天处理失败: {str(e)}")
                yield f"data: {json.dumps({'content': '', 'finished': True, 'error': str(e)})}\n\n"
        
        return StreamingResponse(
            generate_stream(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
                "Access-Control-Allow-Headers": "Content-Type"
            }
        )
        
    except Exception as e:
        logger.error(f"流式聊天端点失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"流式聊天处理失败: {str(e)}")

@router.post("/chat", response_model=ChatResponse)
async def chat_endpoint(request: ChatRequest):
    """
    智能聊天端点
    支持基于RAG的问答和普通聊天
    """
    try:
        logger.info(f"收到聊天请求: {request.message[:50]}...")
        
        if request.use_rag:
            # RAG模式：基于向量存储的文档内容回答
            logger.info("使用RAG模式进行回答")
            
            try:
                # 使用PostgreSQL向量存储服务检索相关文档
                search_results = await postgresql_vector_service.search_similar(
                    query=request.message,
                    top_k=5
                )
                
                if search_results and len(search_results) > 0:
                    # 提取相关文档块
                    relevant_chunks = []
                    for result in search_results:
                        relevant_chunks.append({
                            "content": result.get("content", ""),
                            "metadata": result.get("metadata", {}),
                            "similarity": result.get("similarity", 0.0)
                        })
                    
                    # 生成RAG响应
                    rag_response = await rag_service.generate_response(
                        request.message, relevant_chunks, request.conversation_history
                    )
                    
                    return ChatResponse(
                        response=rag_response["answer"],
                        timestamp=rag_response["timestamp"],
                        token_usage=rag_response.get("token_usage"),
                        sources_used=relevant_chunks
                    )
                else:
                    logger.info("未找到相关文档，使用普通聊天模式")
                    
            except Exception as e:
                logger.error(f"RAG检索失败: {str(e)}")
                logger.info("RAG检索失败，回退到普通聊天模式")
        
        # 普通聊天模式
        logger.info("使用普通聊天模式")
        messages = [
            {"role": "system", "content": "你是一个友好的校园智能助手，可以帮助回答各种问题。"}
        ]
        
        # 添加对话历史
        if request.conversation_history:
            messages.extend(request.conversation_history[-10:])
        
        # 添加当前消息
        messages.append({"role": "user", "content": request.message})
        
        # 调用DeepSeek API
        response = await rag_service.chat_completion(messages)
        
        # 转换token_usage为字典格式
        usage = response.get("usage")
        token_usage_dict = None
        if usage:
            token_usage_dict = {
                "prompt_tokens": usage.prompt_tokens,
                "completion_tokens": usage.completion_tokens,
                "total_tokens": usage.total_tokens
            }
        
        return ChatResponse(
            response=response["content"],
            timestamp=datetime.now().isoformat(),
            token_usage=token_usage_dict,
            sources_used=[]  # 非RAG模式下没有文档来源
        )
        
    except Exception as e:
        logger.error(f"聊天处理失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"聊天处理失败: {str(e)}")

@router.post("/documents/process", response_model=DocumentProcessResponse)
async def process_document_endpoint(request: DocumentProcessRequest):
    """
    文档处理端点
    将文档内容分块、生成嵌入向量并存储到向量数据库
    """
    try:
        logger.info(f"开始处理文档: {request.title or '未命名文档'}")
        
        # 构建元数据
        metadata = request.metadata or {}
        metadata.update({
            "title": request.title,
            "uploaded_at": datetime.now().isoformat()
        })
        
        # 1. 处理文档（分块和向量化）
        processed_doc = await rag_service.process_document(request.content, metadata)
        
        # 2. 存储到向量数据库
        doc_id = f"doc_{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}_{uuid.uuid4().hex[:8]}"
        
        # 准备存储数据
        document_data = {
            "id": doc_id,
            "title": request.title or "未命名文档",
            "content": request.content,
            "chunks": processed_doc["chunks"],
            "embeddings": processed_doc["embeddings"],
            "metadata": metadata,
            "processed_at": processed_doc["processed_at"],
            "chunk_count": processed_doc["chunk_count"]
        }
        
        # 存储到PostgreSQL向量数据库
        await postgresql_vector_service.add_document(
            document_id=doc_id,
            chunks=processed_doc["chunks"],
            metadata={
                "title": request.title or "未命名文档",
                "content": request.content,
                "processed_at": processed_doc["processed_at"],
                "chunk_count": processed_doc["chunk_count"],
                **metadata
            }
        )
        
        # 3. 存储到内存缓存（用于快速访问）
        processed_doc_with_content = {
            **processed_doc,
            "content": request.content,
            "title": request.title or "未命名文档"
        }
        processed_documents[doc_id] = processed_doc_with_content
        
        logger.info(f"文档处理和存储完成，ID: {doc_id}, 块数: {processed_doc['chunk_count']}")
        
        return DocumentProcessResponse(
            success=True,
            document_id=doc_id,
            title=request.title or "未命名文档",
            chunk_count=processed_doc["chunk_count"],
            processing_time=0.0,  # 可以计算实际处理时间
            processed_at=processed_doc["processed_at"],
            message=f"文档处理成功，已生成 {processed_doc['chunk_count']} 个文本块并存储到向量数据库"
        )
        
    except Exception as e:
        logger.error(f"文档处理失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文档处理失败: {str(e)}")

@router.post("/documents/upload")
async def upload_document_endpoint(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    category: Optional[str] = Form("general")
):
    """
    文档上传端点
    支持多种文档格式：PDF、DOCX、TXT、MD等
    现在使用document_service进行完整的文档处理和存储
    """
    try:
        logger.info(f"收到文件上传: {file.filename}, 类型: {file.content_type}")
        
        # 检查文件大小 (最大50MB)
        MAX_FILE_SIZE = 50 * 1024 * 1024
        file.file.seek(0, 2)  # 移动到文件末尾
        file_size = file.file.tell()
        file.file.seek(0)  # 重置到文件开头
        
        if file_size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400, 
                detail=f"文件大小超过限制（最大50MB），当前文件大小: {file_size / (1024*1024):.2f}MB"
            )
        
        if file_size == 0:
            raise HTTPException(status_code=400, detail="文件不能为空")
        
        # 读取文件内容
        content = await file.read()
        
        # 构建元数据
        metadata = {
            "filename": file.filename,
            "content_type": file.content_type,
            "file_size": file_size,
            "category": category,
            "upload_timestamp": datetime.now().isoformat(),
            "title": title or file.filename
        }
        
        # 使用document_service进行完整的文档处理
        # 这将正确地存储到documents表和向量数据库
        result = await document_service.upload_document(
            file_content=content,
            filename=file.filename,
            metadata=metadata
        )
        
        if not result["success"]:
            raise HTTPException(status_code=400, detail=result["error"])
        
        # 构建API响应格式，保持与原来的接口兼容
        response = {
            "success": True,
            "document_id": result["document_id"],  # 现在是正确的整数ID
            "title": title or file.filename,
            "chunk_count": result["chunks_count"],
            "processing_time": 0.0,
            "processed_at": result["timestamp"],
            "message": f"文档上传成功，已生成 {result['chunks_count']} 个文本块并存储到数据库",
            "upload_info": {
                "original_filename": file.filename,
                "file_size_mb": round(file_size / (1024*1024), 2),
                "text_extracted_length": result["content_length"],
                "category": category,
                "file_type": result["file_type"]
            }
        }
        
        logger.info(f"文档上传和处理成功: {file.filename}, ID: {result['document_id']}")
        return response
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文件上传失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文件上传失败: {str(e)}")


async def extract_text_from_file(content: bytes, filename: str, content_type: str) -> str:
    """
    从不同格式的文件中提取文本内容
    """
    try:
        # 获取文件扩展名
        file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
        
        # PDF文件处理
        if content_type == 'application/pdf' or file_extension == 'pdf':
            return await extract_text_from_pdf(content)
        
        # Word文档处理
        elif (content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                              'application/msword'] or file_extension in ['docx', 'doc']):
            return await extract_text_from_docx(content)
        
        # 文本文件处理
        elif (content_type.startswith('text/') or file_extension in ['txt', 'md', 'markdown', 'py', 'js', 'html', 'css', 'json']):
            return await extract_text_from_text(content)
        
        else:
            raise ValueError(f"不支持的文件格式: {content_type} (.{file_extension})")
            
    except Exception as e:
        logger.error(f"文本提取失败: {str(e)}")
        raise


async def extract_text_from_pdf(content: bytes) -> str:
    """从PDF文件提取文本"""
    try:
        import PyPDF2
        import io
        
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_content = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_content += f"\n--- 第{page_num + 1}页 ---\n{page_text}\n"
            except Exception as e:
                logger.warning(f"PDF第{page_num + 1}页文本提取失败: {str(e)}")
                continue
        
        if not text_content.strip():
            raise ValueError("PDF文件中未找到可提取的文本内容")
        
        return text_content.strip()
        
    except ImportError:
        raise ValueError("PDF处理库未安装，请安装PyPDF2: pip install PyPDF2")
    except Exception as e:
        raise ValueError(f"PDF文本提取失败: {str(e)}")


async def extract_text_from_docx(content: bytes) -> str:
    """从DOCX文件提取文本"""
    try:
        from docx import Document
        import io
        
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        
        text_content = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content += " | ".join(row_text) + "\n"
        
        if not text_content.strip():
            raise ValueError("DOCX文件中未找到可提取的文本内容")
        
        return text_content.strip()
        
    except ImportError:
        raise ValueError("DOCX处理库未安装，请安装python-docx: pip install python-docx")
    except Exception as e:
        raise ValueError(f"DOCX文本提取失败: {str(e)}")


async def extract_text_from_text(content: bytes) -> str:
    """从文本文件提取内容"""
    try:
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
        
        for encoding in encodings:
            try:
                text_content = content.decode(encoding)
                return text_content.strip()
            except UnicodeDecodeError:
                continue
        
        raise ValueError("无法识别文件编码，请确保文件使用UTF-8、GBK或其他常见编码")
        
    except Exception as e:
        raise ValueError(f"文本文件处理失败: {str(e)}")

@router.get("/documents")
async def list_documents():
    """
    获取已处理文档列表
    """
    try:
        # 从PostgreSQL向量数据库获取统计信息
        stats = await postgresql_vector_service.get_stats()
        
        # 从PostgreSQL获取文档信息
        documents = await postgresql_vector_service.get_document_chunks()
        
        # 获取documents表中的文档信息
        from app.models.document import Document
        db_documents = Document.get_all(limit=1000)
        
        # 创建文档ID到文档信息的映射
        doc_info_map = {}
        for doc in db_documents:
            doc_info_map[str(doc.id)] = {
                "filename": doc.filename,
                "upload_time": doc.upload_time.isoformat() if doc.upload_time else datetime.now().isoformat()
            }
        
        # 按文档ID分组统计
        document_map = {}
        for chunk in documents:
            doc_id = chunk["document_id"]
            
            if doc_id not in document_map:
                # 优先从documents表获取文件名，如果没有则从chunk metadata获取，最后使用默认值
                title = "未命名文档"
                processed_at = datetime.now().isoformat()
                
                # 确保doc_id是字符串类型进行比较
                doc_id_str = str(doc_id)
                if doc_id_str in doc_info_map:
                    title = doc_info_map[doc_id_str]["filename"]
                    processed_at = doc_info_map[doc_id_str]["upload_time"]
                elif "filename" in chunk["metadata"]:
                    title = chunk["metadata"]["filename"]
                
                document_map[doc_id] = {
                    "id": doc_id,
                    "title": title,
                    "chunk_count": 0,
                    "processed_at": processed_at,
                    "content_length": 0
                }
            
            document_map[doc_id]["chunk_count"] += 1
            document_map[doc_id]["content_length"] += len(chunk["content"])
        
        # 转换为列表
        documents = list(document_map.values())
        
        return {
            "documents": documents,
            "total_count": len(documents)
        }
        
    except Exception as e:
        logger.error(f"获取文档列表失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"获取文档列表失败: {str(e)}")

@router.delete("/documents/{document_id}")
async def delete_document_endpoint(document_id: str):
    """
    删除指定文档
    
    行为说明：
    - 若 `document_id` 可解析为整数，则认为是数据库documents表的主键，调用服务层执行硬删除（含向量、分块、标签和本地文件清理）。
    - 若 `document_id` 非数字（旧格式，如 `doc_XXXX`），则仅删除向量存储中的分块，并清理内存缓存。
    返回删除结果与细节。
    """
    try:
        # 尝试将ID解析为整数，以决定删除策略
        try:
            int_id = int(document_id)
            # 使用服务层执行完整硬删除
            result = await document_service.delete_document(int_id)
            if not result.get("success"):
                raise HTTPException(status_code=400, detail=result.get("error", "删除失败"))

            # 清理内存缓存（新旧ID都尝试）
            if document_id in processed_documents:
                del processed_documents[document_id]
            if str(int_id) in processed_documents:
                del processed_documents[str(int_id)]

            logger.info(f"文档删除成功（硬删除）: {document_id}")
            return {
                "message": "文档删除成功",
                "document_id": document_id,
                "delete_details": result
            }
        except ValueError:
            # 旧格式ID：仅删除向量存储
            delete_result = await postgresql_vector_service.delete_document(document_id)
            # 清理内存缓存
            if document_id in processed_documents:
                del processed_documents[document_id]
            logger.info(f"文档删除成功（仅向量）: {document_id}")
            return {
                "message": "文档向量删除成功",
                "document_id": document_id,
                "delete_details": delete_result
            }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文档删除失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文档删除失败: {str(e)}")

@router.get("/health")
async def health_check_endpoint():
    """
    RAG系统健康检查端点
    """
    try:
        health_status = await rag_service.health_check()
        
        # 添加文档统计信息
        health_status["document_stats"] = {
            "total_documents": len(processed_documents),
            "total_chunks": sum(doc["chunk_count"] for doc in processed_documents.values())
        }
        
        if not health_status["overall"]:
            raise HTTPException(status_code=503, detail="RAG系统服务不可用")
        
        return health_status
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"健康检查失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"健康检查失败: {str(e)}")

@router.post("/search")
async def search_documents_endpoint(request: QueryRequest):
    """
    独立的文档检索API端点
    基于语义相似度搜索相关文档块
    """
    try:
        logger.info(f"收到文档检索请求: {request.query[:50]}...")
        
        start_time = datetime.now()
        
        # 直接使用PostgreSQL向量服务进行搜索
        search_results = await postgresql_vector_service.search_similar(
            query=request.query,
            top_k=request.top_k if hasattr(request, 'top_k') else 10,
            similarity_threshold=request.similarity_threshold if hasattr(request, 'similarity_threshold') else 0.3
        )
        
        # 格式化搜索结果以匹配API响应格式
        formatted_results = []
        for result in search_results:
            metadata = result.get("metadata", {})
            formatted_results.append({
                "content": result["content"],
                "similarity": result["similarity"],
                "document_id": result["document_id"],
                "document_title": metadata.get("title", "未命名文档"),
                "chunk_index": metadata.get("chunk_index", 0),
                "processed_at": metadata.get("processed_at", "")
            })
        
        end_time = datetime.now()
        search_time = (end_time - start_time).total_seconds()
        
        logger.info(f"文档检索完成，找到 {len(formatted_results)} 个相关结果，耗时 {search_time:.3f}秒")
        
        return QueryResponse(
            query=request.query,
            results=formatted_results,
            total_results=len(formatted_results),
            search_time=search_time,
            message=f"成功检索到 {len(formatted_results)} 个相关文档块" if formatted_results else "未找到匹配的文档"
        )
        
    except Exception as e:
        logger.error(f"文档检索失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文档检索失败: {str(e)}")

@router.get("/stats")
async def get_system_stats():
    """
    获取系统统计信息
    """
    try:
        total_chunks = sum(doc["chunk_count"] for doc in processed_documents.values())
        total_content_length = sum(len(doc["content"]) for doc in processed_documents.values())
        
        return {
            "system_stats": {
                "total_documents": len(processed_documents),
                "total_chunks": total_chunks,
                "total_content_length": total_content_length,
                "average_chunks_per_doc": total_chunks / len(processed_documents) if processed_documents else 0,
                "rag_service_config": {
                    "chunk_size": rag_service.chunk_size,
                    "chunk_overlap": rag_service.chunk_overlap,
                    "top_k_retrieval": rag_service.top_k_retrieval,
                    "max_context_length": rag_service.max_context_length
                }
            },
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"获取统计信息失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"获取统计信息失败: {str(e)}")
